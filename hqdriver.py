import errno
import os
from datetime import datetime
import ActionClasses.Browser as b
import ActionClasses.BrowserNav
import ActionClasses.ClickElement
import ActionClasses.CloseBrowser
import ActionClasses.Elements
import ActionClasses.SetText
import ActionClasses.TakeSCreenShot
import pandas as pd
from docx import Document
from docx.shared import Inches


class hqDriver():
    resultDir=None
    df=None
    listScreenShot=[]
    def execute():
        #creating result folder
        hqDriver.resultDir = hqDriver.filecreation()
        #path until UIAF_UI
        dir_path = os.path.dirname(os.path.realpath(__file__))
        #Reading Execution.xlsx
        ExecutionFileName=os.path.dirname(dir_path) + '/Selenium/TestData/Execution.xlsx'
        if ExecutionFileName:
            ecolumnNameChoice = 'TO BE EXECUTED'
            ecolumnNameMasterSheet = 'MASTER_SHEET_EXCEL'
            fields=[ecolumnNameChoice,ecolumnNameMasterSheet]
            df_exec=pd.read_excel(ExecutionFileName,skipinitialspace=True, usecols=fields)
            erowCount=len(df_exec.index)
            for i in range(0,erowCount):
                if df_exec.loc[i][ecolumnNameChoice].lower()=='yes':
                    #Reading master sheet
                    MasterSheetName=os.path.dirname(dir_path) + '/Selenium/TestData/'+df_exec.loc[i][ecolumnNameMasterSheet]
                    if MasterSheetName:
                       mcolumnNameChoice = 'EXECUTABLE'
                       mcolumnNameScriptSheet = 'SCRIPT EXCEL NAME'
                       mfields = [mcolumnNameChoice, mcolumnNameScriptSheet]
                       df_master = pd.read_excel(MasterSheetName, skipinitialspace=True, usecols=mfields,sheet_name='Master')
                       mrowCount = len(df_master.index)
                       for j in range(0,mrowCount):
                           if df_master.loc[j][mcolumnNameChoice].lower()=='yes':
                               #Reading script sheet
                               ScriptSheetName=os.path.dirname(dir_path) + '/Selenium/TestData/'+df_master.loc[i][mcolumnNameScriptSheet]
                               if ScriptSheetName:
                                    sColumnActionHeader='ACTION'
                                    sObjectNameHeader='OBJECTNAME'
                                    sInputData='INPUT DATA'
                                    sScreenShot='SCREENSHOT'
                                    sStepNo='STEP NO'
                                    sTestCaseName='TESTCASE NAME'
                                    sfields = [sColumnActionHeader,sObjectNameHeader,sInputData,sScreenShot,sStepNo,sTestCaseName]
                                    df_script = pd.read_excel(ScriptSheetName, skipinitialspace=True, usecols=sfields)
                                    #Creation of empty dataframe with only column
                                    headerForResultDF=[sTestCaseName,sStepNo,sColumnActionHeader,'Result']
                                    hqDriver.creation_dataframe(headerForResultDF)
                                    oObjectNameHeader='OBJECT VARIABLE NAME'
                                    oObjectPropertyHeader='OBJECT_PROPERTY'
                                    oPropertyValHeader='PROPERTY_VALUE'
                                    ofields=[oObjectNameHeader,oObjectPropertyHeader,oPropertyValHeader]
                                    #Reading all objects sheet
                                    df_script_AllObjectsSheet=pd.read_excel(ScriptSheetName, skipinitialspace=True, usecols=ofields,sheet_name='AllObjects')
                                    srowCount = len(df_script.index)
                                    for k  in range(0,srowCount):
                                      actionName=df_script.loc[k][sColumnActionHeader]
                                      objnameexists=pd.notnull(df_script.loc[k][sObjectNameHeader])
                                      objname=df_script.loc[k][sObjectNameHeader]
                                      stepData=df_script.loc[k][sInputData]
                                      screenShot=df_script.loc[k][sScreenShot]
                                      stepno=df_script.loc[k][sStepNo]
                                      testCaseName=df_script.loc[k][sTestCaseName]
                                      screenShotName = testCaseName+"_"+stepno+"_" + actionName + ".png"
                                      #if there is no object name
                                      if not objnameexists:
                                        result=hqDriver.selectAction(actionName,stepData,None,None)
                                        #ScreenShot Generation
                                        if str(screenShot).lower()=='yes':
                                            s=ActionClasses.TakeSCreenShot.TakeScreenShot.execute(hqDriver.resultDir,screenShotName)
                                            hqDriver.listScreenShot.append(s)
                                        #Writing step result to dataframe
                                        hqDriver.writing_data_to_dataframe(k,testCaseName,stepno,actionName,result)
                                      #if there is object name
                                      else:
                                            orowCount=len(df_script_AllObjectsSheet.index)
                                            #Fetching details of object
                                            for l in range(0,orowCount):
                                                if objname==df_script_AllObjectsSheet.loc[l][oObjectNameHeader]:
                                                    objprop=df_script_AllObjectsSheet.loc[l][oObjectPropertyHeader]
                                                    objval=df_script_AllObjectsSheet.loc[l][oPropertyValHeader]
                                                    result=hqDriver.selectAction(actionName,stepData,objprop,objval)
                                                    #ScreenShot Generation
                                                    if str(screenShot).lower() == 'yes':
                                                       s= ActionClasses.TakeSCreenShot.TakeScreenShot.execute(hqDriver.resultDir,screenShotName)
                                                       hqDriver.listScreenShot.append(s)
                                            # Writing step result to dataframe
                                            hqDriver.writing_data_to_dataframe(k,testCaseName, stepno, actionName, result)
        #Creation of excel report
        hqDriver.excel_test_result_creation()
        #Creation of word doc report
        hqDriver.create_Word_Doc_Result(hqDriver.listScreenShot)
    #all Selenium actions
    def selectAction(stepaction,stepdata,objprop,objval):
        if str(stepaction).lower()=='launchapplication':
            result = b.execute(stepdata)
            return result
        elif str(stepaction).lower()=='settext':
            result=ActionClasses.SetText.SetText.execute(objprop,objval,stepdata)
            return result
        elif str(stepaction).lower() == 'clickelement':
            result = ActionClasses.ClickElement.ClickElement.execute(objprop,objval)
            return result
        elif str(stepaction).lower()=='closebrowser':
            result=ActionClasses.CloseBrowser.CloseBrowser.execute()
            return result
    #Creation of result folder
    def filecreation():
        dir_path = os.path.dirname(os.path.realpath(__file__))
        mydir = os.path.join(os.path.dirname(dir_path),"test results",
            "Result_"+datetime.now().strftime('%Y-%m-%d_%H-%M-%S'))
        try:
            os.makedirs(mydir)
        except OSError as e:
            if e.errno != errno.EEXIST:
                raise  # This was not a "directory exist" error..
        return mydir
    #converting pandas dataframe to excel
    def excel_test_result_creation():
         hqDriver.df.to_excel(hqDriver.resultDir+"\\TestResult.xlsx",index=False)
    #creation of empty dataframe with headers
    def creation_dataframe(headerList):
        hqDriver.df = pd.DataFrame(columns=headerList)
    #Writing data to dataframe
    def writing_data_to_dataframe(row,testCaseName, stepno, actionName,result):
        hqDriver.df.loc[row]=[testCaseName, stepno, actionName,result]
    #Creation of word result
    def create_Word_Doc_Result(screenShotList):
        document = Document()
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text('Screen Shots')
        for i in screenShotList:
          r.add_picture(i)  # r.add_picture('/tmp/foo.jpg')
        document.save(hqDriver.resultDir + '\\' + 'demo.docx')




hqDriver.execute()








